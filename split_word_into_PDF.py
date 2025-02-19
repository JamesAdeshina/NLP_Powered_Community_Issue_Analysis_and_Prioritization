# !pip install PyPDF2 python-docx

import os
from PyPDF2 import PdfReader, PdfWriter
from docx import Document


def split_pdf(input_pdf, output_folder, pages_per_split):
    reader = PdfReader(input_pdf)
    total_pages = len(reader.pages)

    os.makedirs(output_folder, exist_ok=True)

    for i in range(0, total_pages, pages_per_split):
        writer = PdfWriter()
        for j in range(i, min(i + pages_per_split, total_pages)):
            writer.add_page(reader.pages[j])

        output_pdf_path = os.path.join(output_folder, f"split_{i // pages_per_split + 1}.pdf")
        with open(output_pdf_path, "wb") as output_pdf:
            writer.write(output_pdf)

        print(f"Saved: {output_pdf_path}")


def split_word(input_docx, output_folder, paragraphs_per_split):
    doc = Document(input_docx)
    paragraphs = doc.paragraphs
    total_paragraphs = len(paragraphs)

    os.makedirs(output_folder, exist_ok=True)

    for i in range(0, total_paragraphs, paragraphs_per_split):
        new_doc = Document()
        for j in range(i, min(i + paragraphs_per_split, total_paragraphs)):
            new_doc.add_paragraph(paragraphs[j].text)

        output_docx_path = os.path.join(output_folder, f"split_{i // paragraphs_per_split + 1}.docx")
        new_doc.save(output_docx_path)

        print(f"Saved: {output_docx_path}")

# Example usage
split_pdf("example.pdf", "output_pdfs", 5)   # Splits PDF every 5 pages
split_word("Data/community_issues_letters.docx", "output_docs", 12)  # Splits Word file every 10 paragraphs

