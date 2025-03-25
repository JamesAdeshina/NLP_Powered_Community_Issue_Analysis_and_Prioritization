# !pip install PyPDF2 python-docx

import os
import subprocess
from PyPDF2 import PdfReader, PdfWriter
from docx import Document

def convert_docx_to_pdf(input_docx, output_pdf):
    # Use LibreOffice's command-line tool to convert DOCX to PDF
    # Ensure LibreOffice is installed on your Mac first
    subprocess.run([
        'libreoffice', '--headless', '--convert-to', 'pdf',
        '--outdir', os.path.dirname(output_pdf), input_docx
    ])
    print(f"Converted DOCX to PDF: {output_pdf}")


def split_pdf(input_pdf, output_folder, pages_per_split):
    reader = PdfReader(input_pdf)
    total_pages = len(reader.pages)

    os.makedirs(output_folder, exist_ok=True)

    for i in range(0, total_pages, pages_per_split):
        writer = PdfWriter()
        for j in range(i, min(i + pages_per_split, total_pages)):
            writer.add_page(reader.pages[j])

        output_pdf_path = os.path.join(output_folder, f"split_{i // pages_per_split + 1}.pdf")
        with open(output_pdf_path, "wb") as output_pdf:
            writer.write(output_pdf)

        print(f"Saved: {output_pdf_path}")


def split_word(input_docx, output_folder, paragraphs_per_split):
    doc = Document(input_docx)
    paragraphs = doc.paragraphs
    total_paragraphs = len(paragraphs)

    os.makedirs(output_folder, exist_ok=True)

    for i in range(0, total_paragraphs, paragraphs_per_split):
        new_doc = Document()
        for j in range(i, min(i + paragraphs_per_split, total_paragraphs)):
            new_doc.add_paragraph(paragraphs[j].text)

        output_docx_path = os.path.join(output_folder, f"split_{i // paragraphs_per_split + 1}.docx")
        new_doc.save(output_docx_path)

        print(f"Saved: {output_docx_path}")


def split_word_to_pdf_and_pages(input_docx, output_folder, pages_per_split):
    # Step 1: Convert DOCX to PDF using LibreOffice
    pdf_path = os.path.join(output_folder, "document.pdf")
    convert_docx_to_pdf(input_docx, pdf_path)

    # Step 2: Split the PDF
    split_pdf(pdf_path, output_folder, pages_per_split)



# Example usage
split_pdf("example.pdf", "output_pdfs", 5)   # Splits PDF every 5 pages
split_word("Data/community_issues_letters.docx", "output_docs", 12)  # Splits Word file every 10 paragraphs


# Example usage:
split_word_to_pdf_and_pages("Data/community_issues_letters.docx", "output_pdfs", 1)  # Split Word by single pages (after conversion to PDF)

