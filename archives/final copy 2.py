# ------------------ Standard Libraries ------------------
import os
import io
import ssl
import re

# ------------------ # Add to file reading imports ------------------
import PyPDF2
from docx import Document

# ------------------ Data Handling and Numerical Processing ------------------
import pandas as pd
import numpy as np

# ------------------ NLP Libraries (General) ------------------
import nltk
from nltk.tokenize import word_tokenize, sent_tokenize
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
from nltk.sentiment import SentimentIntensityAnalyzer
import emoji
import contractions

# ------------------ Machine Learning and Text Processing ------------------
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.cluster import KMeans
from sklearn.decomposition import LatentDirichletAllocation

# ------------------ Summarization and Topic Extraction ------------------
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lex_rank import LexRankSummarizer
from rake_nltk import Rake

# ------------------ Transformers and Deep Learning Models ------------------
from transformers import pipeline as hf_pipeline
import sentencepiece

# ------------------ Visualization ------------------
import plotly.graph_objects as go

# ------------------ Web App Framework ------------------
import streamlit as st

# ------------------ Fix OpenMP Warning ------------------
os.environ["KMP_DUPLICATE_LIB_OK"] = "True"

# ------------------ SSL Context for NLTK Downloads ------------------
try:
    _create_unverified_https_context = ssl._create_unverified_context
except AttributeError:
    pass
else:
    ssl._create_default_https_context = _create_unverified_https_context


# ------------------ NLTK Downloads ------------------
# (Assuming these have been downloaded already)
# nltk.download('punkt')
# nltk.download('stopwords')
# nltk.download('wordnet')
# nltk.download('vader_lexicon')

# ------------------ Caching Model Loading ------------------
@st.cache_resource
def get_zero_shot_classifier():
    return hf_pipeline("zero-shot-classification", model="facebook/bart-large-mnli")


@st.cache_resource
def get_abstractive_summarizer():
    return hf_pipeline("summarization", model="sshleifer/distilbart-cnn-12-6", revision="a4f8f3e")


@st.cache_resource
def get_sentiment_pipeline():
    return hf_pipeline("sentiment-analysis", model="distilbert-base-uncased-finetuned-sst-2-english")


# ------------------ Paraphrasing Function ------------------
@st.cache_resource
def load_paraphrase_model():
    from transformers import T5Tokenizer, T5ForConditionalGeneration, pipeline as hf_pipeline
    tokenizer = T5Tokenizer.from_pretrained("Vamsi/T5_Paraphrase_Paws", use_fast=False)
    model = T5ForConditionalGeneration.from_pretrained("Vamsi/T5_Paraphrase_Paws")
    paraphrase_pipeline = hf_pipeline(
        "text2text-generation",
        model=model,
        tokenizer=tokenizer,
        max_length=256
    )
    return paraphrase_pipeline


paraphrase_model = load_paraphrase_model()


def paraphrase_text(text):
    input_text = "paraphrase: " + text + " </s>"
    paraphrase = paraphrase_model(input_text, do_sample=False)
    return paraphrase[0]['generated_text']


# ------------------ Helper: Sanitize Text for PDF ------------------
def sanitize_text(text: str) -> str:
    replacements = {
        "\u2014": "-",
        "\u2013": "-",
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u2026": "..."
    }
    for orig, repl in replacements.items():
        text = text.replace(orig, repl)
    return ''.join(c if ord(c) < 256 else '?' for c in text)


# ------------------ Expanded Candidate Labels for Topic Detection ------------------
CANDIDATE_LABELS_TOPIC: tuple[str, ...] = (
    "Waste Management / Public Cleanliness",
    "Water Scarcity",
    "Food Insecurity",
    "Cybersecurity Threats",
    "Delays in NHS Treatment",
    "Underfunded Healthcare Services",
    "Decline in Local Shops / High Street Businesses",
    "High Cost of Living",
    "Overcrowded Public Transport",
    "Homelessness",
    "Lack of Affordable Housing",
    "Noise Pollution",
    "Potholes / Road Maintenance",
    "Traffic Congestion",
    "Air Pollution",
    "School Overcrowding",
    "Crime Rates in Urban Areas",
    "Limited Green Spaces",
    "Aging Infrastructure",
    "Digital Divide",
    "Rising Energy Costs",
    "Housing Quality Issues",
    "Lack of Social Mobility",
    "Climate Change Adaptation",
    "Elderly Care Shortages",
    "Rural Transport Accessibility",
    "Mental Health Service Shortages",
    "Drug and Alcohol Abuse",
    "Gender Pay Gap",
    "Age Discrimination in Employment",
    "Child Poverty",
    "Bureaucratic Delays in Government Services",
    "Lack of Public Restrooms in Urban Areas",
    "Unsafe Cycling Infrastructure",
    "Tackling Modern Slavery",
    "Gentrification and Displacement",
    "Rise in Anti-Social Behaviour",
    "Tackling Fake News and Misinformation",
    "Integration of Immigrant Communities",
    "Parking Problems",
    "Littering in Public Spaces",
    "Speeding Vehicles",
    "Crumbling Pavements",
    "Public Wi-Fi Gaps",
    "Youth Services Cuts",
    "Erosion of Coastal Areas",
    "Flooding in Residential Areas",
    "Loneliness and Social Isolation",
    "Domestic Violence and Abuse",
    "Racial Inequality and Discrimination",
    "LGBTQ+ Rights and Inclusion",
    "Disability Access",
    "Childcare Costs and Availability",
    "Veteran Support",
    "Community Cohesion",
    "Access to Arts and Culture",
    "Biodiversity Loss",
    "Urban Heat Islands",
    "Single-Use Plastics",
    "Education / Skills Development",
    "Community Workshops",
    "Renewable Energy Transition",
    "Food Waste",
    "Deforestation and Land Use",
    "Light Pollution",
    "Soil Degradation",
    "Marine Pollution",
    "Gig Economy Exploitation",
    "Regional Economic Disparities",
    "Skills Shortages",
    "Zero-Hours Contracts",
    "Pension Inequality",
    "Rising Inflation",
    "Small Business Struggles",
    "Post-Brexit Trade Challenges",
    "Automation and Job Loss",
    "Unpaid Internships",
    "Obesity Epidemic",
    "Dental Care Access",
    "Vaccine Hesitancy",
    "Pandemic Preparedness",
    "Nutritional Education",
    "Physical Inactivity",
    "Student Debt",
    "Teacher Shortages",
    "School Funding Cuts",
    "Bullying in Schools",
    "Access to Higher Education",
    "Vocational Training Gaps",
    "Digital Exclusion",
    "Extracurricular Activity Cuts",
    "Aging Public Buildings",
    "Smart City Development",
    "Electric Vehicle Infrastructure",
    "5G Rollout Delays",
    "Flood Defence Upgrades",
    "Rail Network Overcrowding",
    "AI Ethics and Regulation",
    "Space Debris Management",
    "Genetic Engineering Ethics",
    "Climate Migration",
    "Aging Population",
    "Urbanisation Pressures",
    "Data Privacy Concerns",
    "Sustainable Fashion"
)


# ------------------ Topic Functions ------------------
@st.cache_data
def dynamic_topic_label(keywords: str) -> str:
    classifier = get_zero_shot_classifier()
    result = classifier(keywords, CANDIDATE_LABELS_TOPIC)
    best_label = result["labels"][0]
    return best_label


@st.cache_data
def compute_topic(text: str, top_n: int = 5) -> tuple[str, str]:
    rake_extractor = Rake()
    rake_extractor.extract_keywords_from_text(text)
    ranked_phrases = rake_extractor.get_ranked_phrases()
    top_terms = ranked_phrases[:top_n] if len(ranked_phrases) >= top_n else ranked_phrases
    keyword_str = ", ".join(top_terms)
    topic_label = dynamic_topic_label(keyword_str)
    return topic_label, keyword_str


# ------------------ Preprocessing Functions ------------------
@st.cache_data
def remove_email_headers_and_footers(text: str) -> str:
    lines = text.split('\n')
    stripped_lines = [line.strip() for line in lines]
    if "" in stripped_lines:
        first_blank_index = stripped_lines.index("")
        content = "\n".join(lines[first_blank_index + 1:]).strip()
    else:
        content = text
    signature_markers = ('sincerely,', 'regards,', 'best regards,', 'thanks,', 'thank you,')
    final_lines = []
    for line in content.split('\n'):
        if any(line.lower().startswith(marker) for marker in signature_markers):
            break
        final_lines.append(line)
    return "\n".join(final_lines).strip()


def remove_emojis(text):
    return emoji.replace_emoji(text, replace="")


def expand_contractions(text):
    return contractions.fix(text)


def remove_urls(text):
    return re.sub(r'http\S+|www\S+', '', text)


def remove_mentions_hashtags(text):
    text = re.sub(r'@\w+', '', text)
    text = re.sub(r'#\w+', '', text)
    return text


def remove_punctuation(text):
    return re.sub(r'[^\w\s]', '', text)


def remove_numbers(text):
    return re.sub(r'\d+', '', text)


def normalize_repeated_chars(text):
    return re.sub(r'(.)\1{2,}', r'\1', text)


def remove_extra_whitespace(text):
    return re.sub(r'\s+', ' ', text).strip()


def tokenize_and_lower(text):
    tokens = word_tokenize(text)
    return [token.lower() for token in tokens]


def remove_stopwords(tokens):
    stop_words = set(stopwords.words('english'))
    return [token for token in tokens if token not in stop_words]


def lemmatize_tokens(tokens):
    lemmatizer = WordNetLemmatizer()
    return [lemmatizer.lemmatize(token) for token in tokens]


def comprehensive_text_preprocessing(text, use_lemmatization=True):
    text = remove_emojis(text)
    text = expand_contractions(text)
    text = remove_urls(text)
    text = remove_mentions_hashtags(text)
    text = remove_punctuation(text)
    text = remove_numbers(text)
    text = normalize_repeated_chars(text)
    text = remove_extra_whitespace(text)
    tokens = tokenize_and_lower(text)
    tokens = remove_stopwords(tokens)
    if use_lemmatization:
        tokens = lemmatize_tokens(tokens)
    return ' '.join(tokens)


# ------------------ Unsupervised Classification Functions ------------------
def unsupervised_classification(texts, num_clusters=2):
    vectorizer = TfidfVectorizer(stop_words='english', max_features=1000, ngram_range=(1, 2))
    X = vectorizer.fit_transform(texts)
    kmeans = KMeans(n_clusters=num_clusters, random_state=42)
    kmeans.fit(X)
    return kmeans.labels_, vectorizer, kmeans


# ------------------ Dynamic Topic Labeling for Clusters ------------------
candidate_labels = ["Local Problem", "New Initiatives"]


def dynamic_label_clusters(vectorizer, kmeans):
    cluster_labels = {}
    order_centroids = kmeans.cluster_centers_.argsort()[:, ::-1]
    terms = vectorizer.get_feature_names_out()
    for i in range(kmeans.n_clusters):
        top_terms = [terms[ind] for ind in order_centroids[i, :10]]
        keyword_str = ", ".join(top_terms)
        classifier = get_zero_shot_classifier()
        result = classifier(keyword_str, candidate_labels)
        best_label = result["labels"][0]
        cluster_labels[i] = best_label
    return cluster_labels


# ------------------ Topic Modeling ------------------
def topic_modeling(texts, num_topics=1):
    vectorizer = TfidfVectorizer(stop_words='english', max_features=1000, ngram_range=(1, 2))
    X = vectorizer.fit_transform(texts)
    lda = LatentDirichletAllocation(n_components=num_topics, random_state=42, learning_method='batch', max_iter=10)
    lda.fit(X)
    topics = []
    terms = vectorizer.get_feature_names_out()
    for topic_idx, topic in enumerate(lda.components_):
        top_terms = [terms[i] for i in topic.argsort()[:-6:-1]]
        topics.append(", ".join(top_terms))
    return topics


# ------------------ Summarization Functions ------------------
abstractive_summarizer = get_abstractive_summarizer()


def abstractive_summarization(text):
    summary = abstractive_summarizer(text, max_length=50, min_length=25, do_sample=False)
    return summary[0]['summary_text']


def extractive_summarization(text, sentence_count=3):
    parser = PlaintextParser.from_string(text, Tokenizer("english"))
    summarizer = LexRankSummarizer()
    summary = summarizer(parser.document, sentence_count)
    return " ".join([str(sentence) for sentence in summary])


def query_based_summarization(text, query, threshold=0.1, top_n=2):
    sentences = sent_tokenize(text)
    if not sentences:
        return "No relevant information found for the query."
    action_indicators = ["urge", "request", "increase", "arrange", "immediate", "control", "measure", "action",
                         "implement", "improve", "take"]
    is_action_query = any(
        word in query.lower() for word in ["action", "request", "urge", "increase", "immediate", "control", "measure"])
    if is_action_query:
        threshold = 0.05
    corpus = sentences + [query]
    vectorizer_q = TfidfVectorizer().fit(corpus)
    sentence_vectors = vectorizer_q.transform(sentences)
    query_vector = vectorizer_q.transform([query])
    scores = np.dot(sentence_vectors, query_vector.T).toarray().flatten()
    valid_indices = [i for i, score in enumerate(scores) if score >= threshold]
    if not valid_indices:
        return "No relevant information found for the query."
    if is_action_query:
        valid_indices = [i for i in valid_indices if any(kw in sentences[i].lower() for kw in action_indicators)]
        if not valid_indices:
            return "No relevant information found for the query."
    sorted_indices = sorted(valid_indices, key=lambda i: scores[i], reverse=True)[:top_n]
    selected_indices = sorted(sorted_indices)
    summary = " ".join(sentences[i] for i in selected_indices)
    return summary


# ------------------ Personalized Summary Wrapper ------------------
def personalize_summary(summary, summary_type="general"):
    return f" {summary}"


# ------------------ Sentiment Analysis ------------------
sentiment_pipeline = get_sentiment_pipeline()


def sentiment_analysis(text):
    # 1. VADER Analysis
    sia = SentimentIntensityAnalyzer()
    vader_scores = sia.polarity_scores(text)
    vader_compound = vader_scores["compound"]

    # 2. Transformer Analysis
    transformer_result = sentiment_pipeline(text)[0]
    transformer_label = transformer_result["label"]  # "NEGATIVE" or "POSITIVE"
    transformer_score = transformer_result["score"]

    # Debug info is printed to console; remove or comment out if not needed.
    print("VADER compound score:", vader_compound)
    print("Transformer label:", transformer_label)
    print("Transformer score:", transformer_score)

    # 3. Domain-specific override: if negative keywords are found, set sentiment to Negative.
    negative_keywords = ["complaint", "disrupt", "noise", "excessive", "urge", "investigate", "strict", "control",
                         "disturb"]
    if any(keyword in text.lower() for keyword in negative_keywords):
        final_sentiment = "Negative"
    else:
        if transformer_label == "NEGATIVE":
            final_sentiment = "Negative"
        elif transformer_label == "POSITIVE":
            final_sentiment = "Positive"
        else:
            final_sentiment = "Neutral"
        if vader_compound <= -0.3:
            final_sentiment = "Negative"

    explanation = f"The sentiment of the text is {final_sentiment}."
    return {
        "vader_scores": vader_scores,
        "transformer_result": transformer_result,
        "sentiment_label": final_sentiment,
        "explanation": explanation
    }


# ------------------ Aggregated Analysis Plot Functions ------------------
def plot_classification_distribution(class_counts):
    fig = go.Figure([go.Bar(x=class_counts.index, y=class_counts.values)])
    fig.update_layout(title="Classification Distribution", xaxis_title="Category", yaxis_title="Count")
    return fig


def plot_sentiment_distribution(avg_sentiment):
    fig = go.Figure([go.Bar(x=["Average Sentiment Polarity"], y=[avg_sentiment])])
    fig.update_layout(title="Average Sentiment Polarity", xaxis_title="Metric", yaxis_title="Polarity")
    return fig


def plot_sentiment_gauge(polarity):
    fig = go.Figure(go.Indicator(
        mode="gauge+number",
        value=polarity,
        domain={'x': [0, 1], 'y': [0, 1]},
        title={"text": "Sentiment Gauge"},
        gauge={
            "axis": {"range": [-1, 1]},
            "steps": [
                {"range": [-1, -0.3], "color": "red"},
                {"range": [-0.3, 0.3], "color": "yellow"},
                {"range": [0.3, 1], "color": "green"}
            ],
            "threshold": {
                "line": {"color": "black", "width": 4},
                "thickness": 0.75,
                "value": polarity
            }
        }
    ))
    # Add annotations for legend
    fig.update_layout(
        annotations=[
            dict(x=0.15, y=0.1, text="<b>Negative</b>", showarrow=False, font=dict(color="red", size=12)),
            dict(x=0.50, y=0.1, text="<b>Neutral</b>", showarrow=False, font=dict(color="yellow", size=12)),
            dict(x=0.85, y=0.1, text="<b>Positive</b>", showarrow=False, font=dict(color="green", size=12))
        ]
    )
    return fig


# ------------------ Report Generation Functions ------------------
from fpdf import FPDF
from docx import Document


def generate_pdf_report(original_text, abstractive_summary, extractive_summary, query_summary, sentiment_results):
    original_text = sanitize_text(original_text)
    abstractive_summary = sanitize_text(abstractive_summary)
    extractive_summary = sanitize_text(extractive_summary)
    query_summary = sanitize_text(query_summary)
    sentiment_results = sanitize_text(str(sentiment_results))
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(0, 10, txt="Analysis Report", ln=True, align='C')
    pdf.ln(5)
    pdf.multi_cell(0, 10, txt=f"Original Text:\n{original_text}\n")
    pdf.ln(3)
    pdf.multi_cell(0, 10, txt=f"Abstractive Summary:\n{abstractive_summary}\n")
    pdf.ln(3)
    pdf.multi_cell(0, 10, txt=f"Extractive Summary:\n{extractive_summary}\n")
    pdf.ln(3)
    pdf.multi_cell(0, 10, txt=f"Query-based Summary:\n{query_summary}\n")
    pdf.ln(3)
    pdf.multi_cell(0, 10, txt=f"Sentiment Analysis:\n{sentiment_results}\n")
    return pdf.output(dest='S').encode('latin1', errors='replace')


def generate_docx_report(original_text, abstractive_summary, extractive_summary, query_summary, sentiment_results):
    doc = Document()
    doc.add_heading("Analysis Report", level=1)
    doc.add_heading("Original Text", level=2)
    doc.add_paragraph(original_text)
    doc.add_heading("Abstractive Summary", level=2)
    doc.add_paragraph(abstractive_summary)
    doc.add_heading("Extractive Summary", level=2)
    doc.add_paragraph(extractive_summary)
    doc.add_heading("Query-based Summary", level=2)
    doc.add_paragraph(query_summary)
    doc.add_heading("Sentiment Analysis", level=2)
    doc.add_paragraph(str(sentiment_results))
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ------------------ Data Loading ------------------
def load_data(file_path):
    try:
        df = pd.read_csv(file_path, on_bad_lines='skip')
        st.write("Data loaded successfully.")
        return df
    except Exception as e:
        st.write("Error loading data:", e)
        return None


# ------------------ Main UI Pipeline Using Pages ------------------

def get_file_icon_path(num_files, file_extensions):
    """
    Return a path to the correct SVG icon
    given the number of files and the extension set.
    """
    if num_files == 0:
        return "src/img/Empty.svg"

    if num_files == 1:
        # Check if user pasted
        if "paste" in file_extensions:
            return "src/img/Single_default.svg"  # <--- For a single pasted letter
        elif "pdf" in file_extensions and len(file_extensions) == 1:
            return "src/img/Single_Pdf.svg"
        elif "doc" in file_extensions and len(file_extensions) == 1:
            return "src/img/Single_Doc.svg"
        else:
            return "src/img/Single_default.svg"

    # If multiple files
    # If all PDFs
    if file_extensions == {"pdf"}:
        return "src/img/Multiple_Pdf.svg"
    # If all docs
    elif file_extensions == {"doc"}:
        return "src/img/Multiple_Doc.svg"
    # If mixture
    elif "pdf" in file_extensions or "doc" in file_extensions:
        return "src/img/Multiple_Both.svg"
    else:
        return "src/img/Multiple_default.svg"

def pick_sidebar_icon(num_files, file_types):
    if num_files == 0:
        return "src/img/Multiple_Default.svg"
    if num_files == 1:
        ft = next(iter(file_types))
        if ft == "pdf":
            return "src/img/Single_Pdf.svg"
        elif ft == "doc":
            return "src/img/Single_Doc.svg"
        else:
            return "src/img/Single_Default.svg"
    # For multiple files
    if file_types == {"pdf"}:
        return "src/img/Multiple_Pdf.svg"
    elif file_types == {"doc"}:
        return "src/img/Multiple_Doc.svg"
    elif len(file_types) > 1:  # Mixed types
        return "src/img/Multiple_Both.svg"
    else:
        return "src/img/Multiple_Default.svg"

def data_entry_page():
    st.title("Letter Submission (Data Entry)")
    data_mode = st.radio("Choose Input Mode", ["Paste Text", "Upload File"])

    input_text = ""
    uploaded_files = []

    if data_mode == "Paste Text":
        input_text = st.text_area("Paste your letter text here", height=200)
    else:
        uploaded_files = st.file_uploader(
            "Upload files (txt, pdf, doc, docx)",
            type=["txt", "pdf", "doc", "docx"],
            accept_multiple_files=True
        )

    if st.button("Submit"):
        with st.spinner("Processing..."):
            if data_mode == "Paste Text":
                if not input_text.strip():
                    st.warning("Please paste some text before submitting.")
                    return

                st.session_state.input_text = input_text
                st.session_state.data_submitted = True
                st.session_state.data_mode = data_mode
                st.session_state.uploaded_file_info = {
                    "num_files": 1,
                    "file_extensions": {"paste"}
                }
                st.session_state.page = "results"
                st.rerun()

            elif data_mode == "Upload File":
                if not uploaded_files:
                    st.warning("Please upload at least one file.")
                    return

                file_types = []
                extracted_texts = []
                combined_text = ""

                for file in uploaded_files:
                    file_type = file.type
                    file_types.append(file_type)
                    text = ""

                    try:
                        if file_type == "application/pdf":
                            pdf_reader = PyPDF2.PdfReader(file)
                            text = "\n".join([
                                page.extract_text()
                                for page in pdf_reader.pages
                                if page.extract_text()
                            ])

                        elif file_type in ["application/msword",
                                         "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
                            doc = Document(file)
                            text = "\n".join([
                                para.text
                                for para in doc.paragraphs
                                if para.text.strip()
                            ])

                        elif file_type == "text/plain":
                            text = file.getvalue().decode("utf-8")

                    except Exception as e:
                        st.error(f"Error reading {file.name}: {str(e)}")
                        continue

                    if text.strip():
                        extracted_texts.append(text.strip())
                        combined_text += f"\n\n{text.strip()}"

                if not extracted_texts:
                    st.error("Could not extract any text from uploaded files")
                    return

                # Update session state
                st.session_state.uploaded_files_texts = extracted_texts
                st.session_state.input_text = combined_text.strip()
                st.session_state.data_submitted = True
                st.session_state.data_mode = data_mode

                # Handle file type icons (simplified to pdf/doc/other)
                ext_set = set()
                for ft in file_types:
                    ft_lower = ft.lower()
                    if "pdf" in ft_lower:
                        ext_set.add("pdf")
                    elif "msword" in ft_lower or "wordprocessingml.document" in ft_lower:
                        ext_set.add("doc")
                    else:
                        ext_set.add("other")

                st.session_state.uploaded_file_info = {
                    "num_files": len(uploaded_files),
                    "file_extensions": ext_set
                }

                # Route to correct page
                if len(uploaded_files) > 1:
                    st.session_state.page = "aggregated_analysis"
                else:
                    st.session_state.page = "results"

                st.rerun()


def results_page():
    st.title("Individual Letter Analysis")
    if "data_submitted" not in st.session_state or not st.session_state.data_submitted:
        st.warning("No data submitted yet. Please go to the 'Data Entry' page, provide a letter, and click Submit.")
        return

    # Get text based on input method
    if st.session_state.data_mode == "Upload File":
        if "uploaded_files_texts" in st.session_state and len(st.session_state.uploaded_files_texts) >= 1:
            letter_text = st.session_state.uploaded_files_texts[0]
        else:
            st.error("No text found in uploaded file")
            return
    else:
        letter_text = st.session_state.get("input_text", "")

    # Sidebar: Show a dynamic icon based on the upload type.
    # Sidebar: Unified icon handling
    with st.sidebar:
        file_info = st.session_state.get("uploaded_file_info", {})
        num_files = file_info.get("num_files", 0)
        ext_set = file_info.get("file_extensions", set())

        icon_path = pick_sidebar_icon(num_files, ext_set)
        st.image(icon_path, width=150)

        with st.expander("Original Letter", expanded=False):
            st.write(letter_text if letter_text.strip() else "No text available")

    # Continue processing for individual analysis
    letter_clean = comprehensive_text_preprocessing(letter_text)
    classifier = get_zero_shot_classifier()
    if letter_clean.strip():
        classification_result = classifier(letter_clean, candidate_labels)
        letter_class = classification_result["labels"][0]
    else:
        letter_class = "Unclassified (insufficient text)"
        st.warning("Could not classify - extracted text appears empty")
    st.subheader("Classification")
    st.write(f"This letter is classified as: **{letter_class}**")

    topic_label, top_keywords = compute_topic(letter_clean)
    st.subheader("Topic")
    st.write(f"Topic: **{topic_label}**")

    abstractive_res = abstractive_summarization(letter_text)
    extractive_res = extractive_summarization(letter_text)

    st.markdown(
        """
        <style>
        .card {
            box-shadow: 0 4px 8px 0 rgba(0,0,0,0.2);
            transition: 0.3s;
            border-radius: 5px;
            padding: 16px;
            margin: 10px 0;
            background-color: var(--background-color);
            color: var(--text-color);
        }
        .card:hover {
            box-shadow: 0 8px 16px 0 rgba(0,0,0,0.2);
        }
        [data-theme="light"] .card {
            --background-color: #F0F2F6;
            --text-color: #000000;
        }
        [data-theme="dark"] .card {
            --background-color: #262730;
            --text-color: #FFFFFF;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    col1, col2 = st.columns(2)
    with col1:
        st.subheader("💡 Key Takeaways")
        st.markdown(f"""
            <div class="card">
                <p>{abstractive_res}</p>
            </div>
        """, unsafe_allow_html=True)
    with col2:
        st.subheader("🔍 Highlighted Sentences")
        st.markdown(f"""
            <div class="card">
                <p>{extractive_res}</p>
            </div>
        """, unsafe_allow_html=True)

    st.subheader("❓ Inquisitive Summary")
    user_query = st.text_input("Query", "What actions are being urged in the letter?")
    query_res = query_based_summarization(letter_text, query=user_query)
    refined_query_res = paraphrase_text(query_res)
    st.write(personalize_summary(refined_query_res, "query"))

    st.subheader("🗣️ Resident Mood Overview")
    sentiment_results = sentiment_analysis(letter_text)
    vader_compound = sentiment_results["vader_scores"]["compound"]
    sentiment_label = sentiment_results["sentiment_label"]
    explanation = sentiment_results["explanation"]

    col_mood, col_gauge = st.columns(2)
    with col_mood:
        st.write(f"**Mood:** {sentiment_label}")
        st.write(explanation)
    with col_gauge:
        gauge_fig = plot_sentiment_gauge(vader_compound)
        st.plotly_chart(gauge_fig)

    export_format = st.selectbox("Select Export Format", ["PDF", "DOCX", "TXT", "CSV"])
    if export_format == "PDF":
        file_bytes = generate_pdf_report(letter_text, abstractive_res, extractive_res, query_res, sentiment_results)
        st.download_button("Download Report", file_bytes, file_name="analysis_report.pdf", mime="application/pdf")
    elif export_format == "DOCX":
        file_bytes = generate_docx_report(letter_text, abstractive_res, extractive_res, query_res, sentiment_results)
        st.download_button("Download Report", file_bytes, file_name="analysis_report.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    elif export_format == "TXT":
        txt_report = (
            f"Analysis Report\n\nOriginal Text:\n{letter_text}\n\nAbstractive Summary:\n{abstractive_res}\n\n"
            f"Extractive Summary:\n{extractive_res}\n\nQuery-based Summary:\n{query_res}\n\n"
            f"Sentiment Analysis:\n{sentiment_results}"
        )
        st.download_button("Download Report", txt_report, file_name="analysis_report.txt", mime="text/plain")
    elif export_format == "CSV":
        df_report = pd.DataFrame({
            "Original Text": [letter_text],
            "Abstractive Summary": [abstractive_res],
            "Extractive Summary": [extractive_res],
            "Query-based Summary": [query_res],
            "Sentiment Analysis": [str(sentiment_results)]
        })
        csv_report = df_report.to_csv(index=False)
        st.download_button("Download Report", csv_report, file_name="analysis_report.csv", mime="text/csv")

    if st.button("Back to Data Entry"):
        st.session_state.input_text = ""
        st.session_state.data_submitted = False
        st.session_state.page = "data_entry"
        st.rerun()


def aggregated_analysis_page():
    st.title("Comprehensive Letters Analysis")

    # 1) Ensure user has actually uploaded multiple files
    if not st.session_state.get("data_submitted", False):
        st.warning("No data submitted yet. Please go to the 'Data Entry' page and upload multiple files.")
        return

    if "uploaded_files_texts" not in st.session_state or len(st.session_state.uploaded_files_texts) < 2:
        st.warning("No multiple-file data found. Please go to the 'Data Entry' page and upload multiple files.")
        return

    # 2) Show an icon & original letter in the sidebar
    with st.sidebar:
        file_info = st.session_state.get("uploaded_file_info", {})
        num_files = file_info.get("num_files", 0)
        ext_set = file_info.get("file_extensions", set())

        icon_path = pick_sidebar_icon(num_files, ext_set)
        st.image(icon_path, width=150)

        with st.expander("Original Letter", expanded=False):
            st.write(st.session_state.get("input_text", "No text available."))

    # 3) Prepare data
    uploaded_texts = st.session_state.uploaded_files_texts
    df_agg = pd.DataFrame({"text": uploaded_texts})
    df_agg["clean_text"] = df_agg["text"].apply(comprehensive_text_preprocessing)
    texts_clean = df_agg["clean_text"].tolist()

    # 4) Classification
    labels, vectorizer, kmeans = unsupervised_classification(texts_clean, num_clusters=2)
    cluster_mapping = dynamic_label_clusters(vectorizer, kmeans)
    df_agg["classification"] = [cluster_mapping[label] for label in labels]

    st.subheader("Classification Distribution")
    class_counts = df_agg["classification"].value_counts()
    st.write(class_counts)
    st.plotly_chart(plot_classification_distribution(class_counts))

    # 5) Topic Modeling for each classification
    for category in candidate_labels:
        subset_texts = df_agg[df_agg["classification"] == category]["clean_text"].tolist()
        if subset_texts:
            topics = topic_modeling(subset_texts, num_topics=5)
            st.subheader(f"Extracted Topics for {category}")
            for topic in topics:
                dynamic_label = dynamic_topic_label(topic)
                st.write(f"{dynamic_label} (Keywords: {topic})")

    # 6) Aggregated sentiment
    def get_vader_compound(txt):
        return sentiment_analysis(txt)["vader_scores"]["compound"]

    df_agg["sentiment_polarity"] = df_agg["text"].apply(get_vader_compound)
    st.subheader("Average Sentiment Polarity")
    avg_sentiment = df_agg["sentiment_polarity"].mean()
    st.write(avg_sentiment)
    st.plotly_chart(plot_sentiment_distribution(avg_sentiment))

    st.subheader("Sentiment Gauge (Aggregated)")
    st.plotly_chart(plot_sentiment_gauge(avg_sentiment))

    # 7) Export options
    report_csv = df_agg.to_csv(index=False)
    st.download_button("Download Report (CSV)", report_csv, file_name="aggregated_report.csv", mime="text/csv")

    pdf_bytes = generate_pdf_report("Aggregated Report", "N/A", "N/A", "N/A", df_agg.to_dict())
    st.download_button("Download Report (PDF)", pdf_bytes, file_name="aggregated_report.pdf", mime="application/pdf")

    docx_bytes = generate_docx_report("Aggregated Report", "N/A", "N/A", "N/A", df_agg.to_dict())
    st.download_button("Download Report (DOCX)", docx_bytes, file_name="aggregated_report.docx",
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    # 8) Navigation
    if st.button("Back to Data Entry"):
        st.session_state.input_text = ""
        st.session_state.data_submitted = False
        st.session_state.page = "data_entry"
        st.rerun()


def main():
    st.set_page_config(layout="wide", page_title="Bolsover District Council - Analysis")
    st.sidebar.image("src/img/Bolsover_District_Council_logo.svg", width=150)

    # Only if data has been submitted in Upload File mode, show dynamic icon and original letter expander

    # Page selection logic
    if "page" not in st.session_state:
        st.session_state.page = "data_entry"

    if st.session_state.page == "data_entry":
        data_entry_page()
    elif st.session_state.page == "results":
        results_page()
    elif st.session_state.page == "aggregated_analysis":
        aggregated_analysis_page()


if __name__ == '__main__':
    main()
