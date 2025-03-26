import io
import PyPDF2
import pandas as pd
from fpdf import FPDF
from docx import Document
from config import UK_POSTCODE_REGEX  # Only if needed


def extract_text_from_file(file):
    file_type = file.type
    text = ""

    try:
        if file_type == "application/pdf":
            pdf_reader = PyPDF2.PdfReader(file)
            text = "\n".join([
                page.extract_text()
                for page in pdf_reader.pages
                if page.extract_text()
            ])
        elif file_type in ["application/msword",
                           "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
            doc = Document(file)
            text = "\n".join([
                para.text
                for para in doc.paragraphs
                if para.text.strip()
            ])
        elif file_type == "text/plain":
            text = file.getvalue().decode("utf-8")
    except Exception as e:
        raise ValueError(f"Error reading file: {str(e)}")

    return text


def generate_pdf_report(original_text, abstractive_summary, extractive_summary, query_summary, sentiment_results):
    from fpdf import FPDF

    def sanitize_text(text: str) -> str:
        replacements = {
            "\u2014": "-",
            "\u2013": "-",
            "\u2018": "'",
            "\u2019": "'",
            "\u201c": '"',
            "\u201d": '"',
            "\u2026": "..."
        }
        for orig, repl in replacements.items():
            text = text.replace(orig, repl)
        return ''.join(c if ord(c) < 256 else '?' for c in text)

    original_text = sanitize_text(original_text)
    abstractive_summary = sanitize_text(abstractive_summary)
    extractive_summary = sanitize_text(extractive_summary)
    query_summary = sanitize_text(query_summary)
    sentiment_results = sanitize_text(str(sentiment_results))

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(0, 10, txt="Analysis Report", ln=True, align='C')
    pdf.ln(5)
    pdf.multi_cell(0, 10, txt=f"Original Text:\n{original_text}\n")
    pdf.ln(3)
    pdf.multi_cell(0, 10, txt=f"Abstractive Summary:\n{abstractive_summary}\n")
    pdf.ln(3)
    pdf.multi_cell(0, 10, txt=f"Extractive Summary:\n{extractive_summary}\n")
    pdf.ln(3)
    pdf.multi_cell(0, 10, txt=f"Query-based Summary:\n{query_summary}\n")
    pdf.ln(3)
    pdf.multi_cell(0, 10, txt=f"Sentiment Analysis:\n{sentiment_results}\n")
    return pdf.output(dest='S').encode('latin1', errors='replace')


def generate_docx_report(original_text, abstractive_summary, extractive_summary, query_summary, sentiment_results):
    from docx import Document
    doc = Document()
    doc.add_heading("Analysis Report", level=1)
    doc.add_heading("Original Text", level=2)
    doc.add_paragraph(original_text)
    doc.add_heading("Abstractive Summary", level=2)
    doc.add_paragraph(abstractive_summary)
    doc.add_heading("Extractive Summary", level=2)
    doc.add_paragraph(extractive_summary)
    doc.add_heading("Query-based Summary", level=2)
    doc.add_paragraph(query_summary)
    doc.add_heading("Sentiment Analysis", level=2)
    doc.add_paragraph(str(sentiment_results))
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()