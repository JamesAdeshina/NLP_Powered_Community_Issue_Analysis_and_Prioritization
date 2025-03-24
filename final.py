# ------------------ Standard Libraries ------------------
import os
import io
import ssl
import re

# ------------------ File Handling Libraries ------------------
import PyPDF2
from docx import Document

# ------------------ Data Handling and Numerical Processing ------------------
import pandas as pd
import numpy as np

# ------------------ NLP Libraries ------------------
import nltk
from nltk.tokenize import word_tokenize, sent_tokenize
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
from nltk.sentiment import SentimentIntensityAnalyzer
import emoji
import contractions

# ------------------ Machine Learning and Text Processing ------------------
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.cluster import KMeans
from sklearn.decomposition import LatentDirichletAllocation
import torch

# ------------------ Summarization and Topic Extraction ------------------
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lex_rank import LexRankSummarizer
from rake_nltk import Rake

# ------------------ Transformers and Deep Learning Models ------------------
from transformers import pipeline
from sentence_transformers import SentenceTransformer, util
from transformers import AutoTokenizer, AutoModelForTokenClassification

# ------------------ Map Visualization ------------------
from geopy.geocoders import Nominatim
from opencage.geocoder import OpenCageGeocode
import pydeck as pdk
import requests

# ------------------ Visualization ------------------
import plotly.graph_objects as go
import plotly.express as px

# ------------------ Web App Framework ------------------
import streamlit as st

# ------------------ Fix OpenMP Warning ------------------
os.environ["KMP_DUPLICATE_LIB_OK"] = "True"

# ------------------ SSL Context for NLTK Downloads ------------------
try:
    _create_unverified_https_context = ssl._create_unverified_context
except AttributeError:
    pass
else:
    ssl._create_default_https_context = _create_unverified_https_context


# ------------------ NLTK Downloads ------------------
# (Assuming these have been downloaded already)
# nltk.download('punkt')
# nltk.download('stopwords')
# nltk.download('wordnet')
# nltk.download('vader_lexicon')

# ------------------ Caching Model Loading ------------------
@st.cache_resource
def get_zero_shot_classifier():
    return pipeline(
        "zero-shot-classification",
        model="facebook/bart-large-mnli",
        tokenizer="facebook/bart-large-mnli",  # Explicit tokenizer specification
        framework="pt"  # Specify PyTorch
    )

@st.cache_resource
def get_abstractive_summarizer():
    return pipeline(
        "summarization",
        model="sshleifer/distilbart-cnn-12-6",
        revision="a4f8f3e"
    )

@st.cache_resource
def get_sentiment_pipeline():
    return pipeline(
        "sentiment-analysis",
        model="distilbert-base-uncased-finetuned-sst-2-english",
        return_all_scores=True
    )


@st.cache_resource
def get_qa_pipeline():
    return pipeline(
        "question-answering",
        model="deepset/roberta-base-squad2",
        tokenizer="deepset/roberta-base-squad2"
    )

@st.cache_resource
def get_embeddings_model():
    return SentenceTransformer('sentence-transformers/all-MiniLM-L6-v2')


# ------------------ Location Extraction Functions ------------------
@st.cache_resource
def get_ner_pipeline():
    tokenizer = AutoTokenizer.from_pretrained("dslim/bert-base-NER")
    model = AutoModelForTokenClassification.from_pretrained("dslim/bert-base-NER")
    return hf_pipeline("ner", model=model, tokenizer=tokenizer)

UK_POSTCODE_REGEX = r'\b[A-Z]{1,2}\d{1,2}[A-Z]?\s\d[A-Z]{2}\b'
UK_ADDRESS_REGEX = r'\b\d+\s[\w\s]+\b,\s[\w\s]+,\s' + UK_POSTCODE_REGEX

# Function to geocode addresses

def resolve_postcode_to_address(postcode, api_key):
    """
    Resolves a postcode to a full address using OpenCage Geocoder.
    """
    geocoder = OpenCageGeocode(api_key)

    try:
        # Query the geocoder with the postcode
        results = geocoder.geocode(f"{postcode}, UK")

        if results and len(results) > 0:
            # Extract the formatted address from the first result
            full_address = results[0]['formatted']
            return full_address
        else:
            st.warning(f"Could not resolve postcode {postcode} to a full address.")
            return None
    except Exception as e:
        st.error(f"Error resolving postcode {postcode}: {str(e)}")
        return None



def geocode_addresses(addresses):
    """
    Geocodes a list of addresses or postcodes using OpenCage Geocoder.
    """
    # api_key = st.secrets["e760785d8c7944888beefc24aa42eb66"]
    api_key = "e760785d8c7944888beefc24aa42eb66"  # Replace with your OpenCage API key
    geocoder = OpenCageGeocode(api_key)
    locations = []

    for address in addresses:
        if not address:  # Skip None or empty addresses
            locations.append((None, None))
            continue

        try:
            # If the address is a postcode (e.g., "S44 6JJ"), resolve it to a full address
            if re.match(r'^[A-Za-z]{1,2}\d{1,2}[A-Za-z]?\s*\d[A-Za-z]{2}$', address.strip()):
                full_address = resolve_postcode_to_address(address, api_key)
                if not full_address:
                    st.warning(f"Skipping postcode {address} (could not resolve to full address).")
                    locations.append((None, None))
                    continue
            else:
                full_address = f"{address}, Bolsover, UK"

            # Geocode the full address
            results = geocoder.geocode(full_address)

            if results and len(results) > 0:
                lat = results[0]['geometry']['lat']
                lon = results[0]['geometry']['lng']
                locations.append((lat, lon))
                # st.write(f"Geocoded {full_address}: ({lat}, {lon})")  # Debug output
            else:
                st.warning(f"Geocoding failed for address: {full_address}")
                locations.append((None, None))
        except Exception as e:
            st.error(f"Geocoding error for {address}: {str(e)}")
            locations.append((None, None))

    return locations

# Function to create the clustered map using PyDeck
def create_clustered_map(df, filter_by_sentiment=None, filter_by_issue=None, filter_by_topic=None):
    # Apply filters
    if filter_by_sentiment:
        df = df[df["sentiment"] == filter_by_sentiment]
    if filter_by_issue:
        df = df[df["Issue"] == filter_by_issue]
    if filter_by_topic:
        df = df[df["Topic"] == filter_by_topic]

    # Drop rows where geocoding failed
    df = df.dropna(subset=['lat', 'lon'])

    if df.empty:
        st.error("No data found for the selected filters.")
        return None

    # Define color mapping for sentiment
    color_mapping = {
        "POSITIVE": [0, 0, 255, 160],  # Blue for positive
        "NEGATIVE": [255, 0, 0, 160],   # Red for negative
    }

    # Create a ScatterplotLayer
    layer = pdk.Layer(
        "ScatterplotLayer",
        df,
        get_position=["lon", "lat"],
        get_color="[color_mapping[sentiment][0], color_mapping[sentiment][1], color_mapping[sentiment][2], color_mapping[sentiment][3]]",
        get_radius=100,  # Radius of the points
        pickable=True,
    )

    # Set the viewport location
    view_state = pdk.ViewState(
        longitude=df["lon"].mean(),
        latitude=df["lat"].mean(),
        zoom=12,
        pitch=0,  # No tilt for now
        bearing=0,
    )

    # Create the PyDeck deck
    deck = pdk.Deck(
        layers=[layer],
        initial_view_state=view_state,
        map_style="mapbox://styles/mapbox/light-v10",
        tooltip={
            "html": "<b>Issue:</b> {Issue}<br><b>Sentiment:</b> {sentiment}<br><b>Topic:</b> {Topic}<br><b>Address:</b> {Address}<br><b>Text:</b> {text}",
            "style": {
                "backgroundColor": "steelblue",
                "color": "white",
            },
        },
    )

    return deck


# Function to extract addresses, issues, and topics from text
def extract_addresses_issues_and_topics(text):
    # Extract addresses using the existing UK address regex
    addresses = re.findall(UK_ADDRESS_REGEX, text, flags=re.IGNORECASE)

    # Extract issues using topic detection
    issue, _ = compute_topic(text)  # Use the existing compute_topic function

    # Extract topics
    topic, _ = compute_topic(text)  # Use the existing compute_topic function
    return addresses, issue, topic


def extract_locations(text):
    # Combined approach: NER + UK-specific regex patterns
    locations = set()

    # 1. Extract using NER model
    ner_pipeline = get_ner_pipeline()
    entities = ner_pipeline(text)
    locations.update(entity['word'] for entity in entities if entity['entity'] in ['B-LOC', 'I-LOC'])

    # 2. Find full addresses using regex
    addresses = re.findall(UK_ADDRESS_REGEX, text, flags=re.IGNORECASE)
    locations.update(addresses)

    # 3. Extract postcodes separately
    postcodes = re.findall(UK_POSTCODE_REGEX, text, flags=re.IGNORECASE)
    locations.update(postcodes)

    # Return the first valid address or postcode as a string
    for loc in locations:
        if re.match(UK_ADDRESS_REGEX, loc, flags=re.IGNORECASE) or re.match(UK_POSTCODE_REGEX, loc, flags=re.IGNORECASE):
            return loc  # Return the first valid address or postcode

    return None  # Return None if no valid address or postcode is found

# ------------------ Updated Geocoding Function ------------------
@st.cache_data
def geocode_location(location_name):
    geolocator = Nominatim(user_agent="bolsover_analysis")

    # Try postcode first for better accuracy
    postcode_match = re.search(UK_POSTCODE_REGEX, location_name, re.IGNORECASE)
    if postcode_match:
        try:
            location = geolocator.geocode(postcode_match.group(0), exactly_one=True)
            if location:
                return (location.latitude, location.longitude)
        except Exception as e:
            print(f"Postcode geocoding error: {str(e)}")

    # Fallback to full address
    try:
        location = geolocator.geocode(location_name + ", Bolsover District, UK", exactly_one=True)
        if location:
            return (location.latitude, location.longitude)
    except Exception as e:
        print(f"Address geocoding error: {str(e)}")

    return (None, None)


# Function to create the map
def create_bolsover_map(df):
    # Geocode addresses
    df[['lat', 'lon']] = pd.DataFrame(geocode_addresses(df['Address']), columns=['lat', 'lon'])

    # Drop rows where geocoding failed
    df = df.dropna(subset=['lat', 'lon'])

    # Create the map
    fig = px.scatter_mapbox(
        df,
        lat="lat",
        lon="lon",
        color="Issue",  # Color by issue type
        hover_name="Address",  # Show address on hover
        hover_data=["text"],  # Show letter text on hover
        zoom=12,  # Adjust zoom level for Bolsover
        height=600,  # Map height
        title="Geographical Distribution of Issues in Bolsover"
    )

    # Update map layout
    fig.update_layout(
        mapbox_style="open-street-map",  # Use OpenStreetMap style
        margin={"r": 0, "t": 40, "l": 0, "b": 0},  # Adjust margins
        legend=dict(
            orientation="h",  # Horizontal legend
            yanchor="bottom",
            y=1.02,  # Position legend above the map
            xanchor="right",
            x=1
        )
    )

    return fig


# ------------------ Paraphrasing Function ------------------
@st.cache_resource
def load_paraphrase_model():
    from transformers import T5Tokenizer, T5ForConditionalGeneration, pipeline as hf_pipeline
    tokenizer = T5Tokenizer.from_pretrained("Vamsi/T5_Paraphrase_Paws", use_fast=False)
    model = T5ForConditionalGeneration.from_pretrained("Vamsi/T5_Paraphrase_Paws")
    paraphrase_pipeline = hf_pipeline(
        "text2text-generation",
        model=model,
        tokenizer=tokenizer,
        max_length=256
    )
    return paraphrase_pipeline


paraphrase_model = load_paraphrase_model()


def paraphrase_text(text):
    input_text = "paraphrase: " + text + " </s>"
    paraphrase = paraphrase_model(input_text, do_sample=False)
    return paraphrase[0]['generated_text']


# ------------------ Helper: Sanitize Text for PDF ------------------
def sanitize_text(text: str) -> str:
    replacements = {
        "\u2014": "-",
        "\u2013": "-",
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u2026": "..."
    }
    for orig, repl in replacements.items():
        text = text.replace(orig, repl)
    return ''.join(c if ord(c) < 256 else '?' for c in text)


# ------------------ AI Search Implementation ------------------
def ai_question_answer(question, documents):
    # 1. Document Chunking
    tokenizer = AutoTokenizer.from_pretrained("deepset/roberta-base-squad2")
    max_length = 384  # Model's max input size
    chunks = []
    for doc in documents:
        words = doc.split()
        current_chunk = []
        current_length = 0
        for word in words:
            current_chunk.append(word)
            current_length += len(tokenizer.tokenize(word))
            if current_length >= max_length - 2:  # Account for [CLS] and [SEP]
                chunks.append(" ".join(current_chunk))
                current_chunk = []
                current_length = 0
        if current_chunk:
            chunks.append(" ".join(current_chunk))

    # 2. Semantic Search
    embedder = get_embeddings_model()
    question_embedding = embedder.encode(question, convert_to_tensor=True)
    doc_embeddings = embedder.encode(chunks, convert_to_tensor=True)

    # Find top 3 relevant chunks
    cos_scores = util.cos_sim(question_embedding, doc_embeddings)[0]
    top_results = torch.topk(cos_scores, k=min(3, len(chunks)))

    # 3. Answer Generation
    qa_pipeline = get_qa_pipeline()
    context = " ".join([chunks[i] for i in top_results.indices])

    # Handle no context case
    if not context.strip():
        return "I couldn't find relevant information in the documents to answer that question."

    # Generate answer
    result = qa_pipeline(
        question=question,
        context=context,
        max_answer_len=150,
        handle_impossible_answer=True
    )

    # Format response
    if result['score'] < 0.1:  # Confidence threshold
        return "I'm not entirely sure, but based on the documents: " + result['answer']

    return result['answer']


# ------------------ Updated Data Processing ------------------
def process_uploaded_data(_uploaded_files_texts):
    df = pd.DataFrame({"text": _uploaded_files_texts})

    # Existing processing
    df["clean_text"] = df["text"].apply(comprehensive_text_preprocessing)
    df["classification"] = df["clean_text"].apply(classify_document)
    df["sentiment"] = df["text"].apply(lambda x: sentiment_analysis(x)["sentiment_label"])

    # Enhanced location processing
    df["locations"] = df["text"].apply(
        lambda t: re.findall(UK_ADDRESS_REGEX, t, flags=re.IGNORECASE)
    )
    df["postcodes"] = df["text"].apply(
        lambda t: re.findall(UK_POSTCODE_REGEX, t, flags=re.IGNORECASE)
    )

    # Combine addresses and postcodes
    df["all_locations"] = df.apply(
        lambda row: list(set(row["locations"] + row["postcodes"])),
        axis=1
    )

    # Geocode all found locations
    df["geocoded"] = df["all_locations"].apply(
        lambda locs: [geocode_location(loc) for loc in locs if loc]
    )

    # Explode locations into individual rows
    df = df.explode("geocoded").reset_index(drop=True)
    df[["lat", "lon"]] = pd.DataFrame(
        df["geocoded"].tolist(),
        index=df.index
    )

    return df.dropna(subset=["lat", "lon"])



def classify_document(text):
    classifier = get_zero_shot_classifier()
    result = classifier(text, ["Local Problem", "New Initiative"])
    return result["labels"][0]


# ------------------ Expanded Candidate Labels for Topic Detection ------------------
CANDIDATE_LABELS_TOPIC: tuple[str, ...] = (
    "Waste Management / Public Cleanliness",
    "Water Scarcity",
    "Food Insecurity",
    "Cybersecurity Threats",
    "Delays in NHS Treatment",
    "Underfunded Healthcare Services",
    "Decline in Local Shops / High Street Businesses",
    "High Cost of Living",
    "Overcrowded Public Transport",
    "Homelessness",
    "Lack of Affordable Housing",
    "Noise Pollution",
    "Potholes / Road Maintenance",
    "Traffic Congestion",
    "Air Pollution",
    "School Overcrowding",
    "Crime Rates in Urban Areas",
    "Limited Green Spaces",
    "Aging Infrastructure",
    "Digital Divide",
    "Rising Energy Costs",
    "Housing Quality Issues",
    "Lack of Social Mobility",
    "Climate Change Adaptation",
    "Elderly Care Shortages",
    "Rural Transport Accessibility",
    "Mental Health Service Shortages",
    "Drug and Alcohol Abuse",
    "Gender Pay Gap",
    "Age Discrimination in Employment",
    "Child Poverty",
    "Bureaucratic Delays in Government Services",
    "Lack of Public Restrooms in Urban Areas",
    "Unsafe Cycling Infrastructure",
    "Tackling Modern Slavery",
    "Gentrification and Displacement",
    "Rise in Anti-Social Behaviour",
    "Tackling Fake News and Misinformation",
    "Integration of Immigrant Communities",
    "Parking Problems",
    "Littering in Public Spaces",
    "Speeding Vehicles",
    "Crumbling Pavements",
    "Public Wi-Fi Gaps",
    "Youth Services Cuts",
    "Erosion of Coastal Areas",
    "Flooding in Residential Areas",
    "Loneliness and Social Isolation",
    "Domestic Violence and Abuse",
    "Racial Inequality and Discrimination",
    "LGBTQ+ Rights and Inclusion",
    "Disability Access",
    "Childcare Costs and Availability",
    "Veteran Support",
    "Community Cohesion",
    "Access to Arts and Culture",
    "Biodiversity Loss",
    "Urban Heat Islands",
    "Single-Use Plastics",
    "Education / Skills Development",
    "Community Workshops",
    "Renewable Energy Transition",
    "Food Waste",
    "Deforestation and Land Use",
    "Light Pollution",
    "Soil Degradation",
    "Marine Pollution",
    "Gig Economy Exploitation",
    "Regional Economic Disparities",
    "Skills Shortages",
    "Zero-Hours Contracts",
    "Pension Inequality",
    "Rising Inflation",
    "Small Business Struggles",
    "Post-Brexit Trade Challenges",
    "Automation and Job Loss",
    "Unpaid Internships",
    "Obesity Epidemic",
    "Dental Care Access",
    "Vaccine Hesitancy",
    "Pandemic Preparedness",
    "Nutritional Education",
    "Physical Inactivity",
    "Student Debt",
    "Teacher Shortages",
    "School Funding Cuts",
    "Bullying in Schools",
    "Access to Higher Education",
    "Vocational Training Gaps",
    "Digital Exclusion",
    "Extracurricular Activity Cuts",
    "Aging Public Buildings",
    "Smart City Development",
    "Electric Vehicle Infrastructure",
    "5G Rollout Delays",
    "Flood Defence Upgrades",
    "Rail Network Overcrowding",
    "AI Ethics and Regulation",
    "Space Debris Management",
    "Genetic Engineering Ethics",
    "Climate Migration",
    "Aging Population",
    "Urbanisation Pressures",
    "Data Privacy Concerns",
    "Sustainable Fashion"
)


# ------------------ Topic Functions ------------------
@st.cache_data
def dynamic_topic_label(keywords: str) -> str:
    classifier = get_zero_shot_classifier()
    result = classifier(keywords, CANDIDATE_LABELS_TOPIC)
    best_label = result["labels"][0]
    return best_label


@st.cache_data
def compute_topic(text: str, top_n: int = 5) -> tuple[str, str]:
    rake_extractor = Rake()
    rake_extractor.extract_keywords_from_text(text)
    ranked_phrases = rake_extractor.get_ranked_phrases()
    top_terms = ranked_phrases[:top_n] if len(ranked_phrases) >= top_n else ranked_phrases
    keyword_str = ", ".join(top_terms)
    topic_label = dynamic_topic_label(keyword_str)
    return topic_label, keyword_str


# ------------------ Preprocessing Functions ------------------
@st.cache_data
def remove_email_headers_and_footers(text: str) -> str:
    lines = text.split('\n')
    stripped_lines = [line.strip() for line in lines]
    if "" in stripped_lines:
        first_blank_index = stripped_lines.index("")
        content = "\n".join(lines[first_blank_index + 1:]).strip()
    else:
        content = text
    signature_markers = ('sincerely,', 'regards,', 'best regards,', 'thanks,', 'thank you,')
    final_lines = []
    for line in content.split('\n'):
        if any(line.lower().startswith(marker) for marker in signature_markers):
            break
        final_lines.append(line)
    return "\n".join(final_lines).strip()


def remove_emojis(text):
    return emoji.replace_emoji(text, replace="")


def expand_contractions(text):
    return contractions.fix(text)


def remove_urls(text):
    return re.sub(r'http\S+|www\S+', '', text)


def remove_mentions_hashtags(text):
    text = re.sub(r'@\w+', '', text)
    text = re.sub(r'#\w+', '', text)
    return text


def remove_punctuation(text):
    return re.sub(r'[^\w\s]', '', text)


def remove_numbers(text):
    return re.sub(r'\d+', '', text)


def normalize_repeated_chars(text):
    return re.sub(r'(.)\1{2,}', r'\1', text)


def remove_extra_whitespace(text):
    return re.sub(r'\s+', ' ', text).strip()


def tokenize_and_lower(text):
    tokens = word_tokenize(text)
    return [token.lower() for token in tokens]


def remove_stopwords(tokens):
    stop_words = set(stopwords.words('english'))
    return [token for token in tokens if token not in stop_words]


def lemmatize_tokens(tokens):
    lemmatizer = WordNetLemmatizer()
    return [lemmatizer.lemmatize(token) for token in tokens]


def comprehensive_text_preprocessing(text, use_lemmatization=True):
    text = remove_emojis(text)
    text = expand_contractions(text)
    text = remove_urls(text)
    text = remove_mentions_hashtags(text)
    text = remove_punctuation(text)
    text = remove_numbers(text)
    text = normalize_repeated_chars(text)
    text = remove_extra_whitespace(text)
    tokens = tokenize_and_lower(text)
    tokens = remove_stopwords(tokens)
    if use_lemmatization:
        tokens = lemmatize_tokens(tokens)
    return ' '.join(tokens)


# ------------------ Unsupervised Classification Functions ------------------
def unsupervised_classification(texts, num_clusters=2):
    vectorizer = TfidfVectorizer(stop_words='english', max_features=1000, ngram_range=(1, 2))
    X = vectorizer.fit_transform(texts)
    kmeans = KMeans(n_clusters=num_clusters, random_state=42)
    kmeans.fit(X)
    return kmeans.labels_, vectorizer, kmeans


# ------------------ Dynamic Topic Labeling for Clusters ------------------
candidate_labels = ["Local Problem", "New Initiatives"]


def dynamic_label_clusters(vectorizer, kmeans):
    cluster_labels = {}
    order_centroids = kmeans.cluster_centers_.argsort()[:, ::-1]
    terms = vectorizer.get_feature_names_out()
    for i in range(kmeans.n_clusters):
        top_terms = [terms[ind] for ind in order_centroids[i, :10]]
        keyword_str = ", ".join(top_terms)
        classifier = get_zero_shot_classifier()
        result = classifier(keyword_str, candidate_labels)
        best_label = result["labels"][0]
        cluster_labels[i] = best_label
    return cluster_labels


# ------------------ Topic Modeling ------------------
def topic_modeling(texts, num_topics=1):
    vectorizer = TfidfVectorizer(stop_words='english', max_features=1000, ngram_range=(1, 2))
    X = vectorizer.fit_transform(texts)
    lda = LatentDirichletAllocation(n_components=num_topics, random_state=42, learning_method='batch', max_iter=10)
    lda.fit(X)
    topics = []
    terms = vectorizer.get_feature_names_out()
    for topic_idx, topic in enumerate(lda.components_):
        top_terms = [terms[i] for i in topic.argsort()[:-6:-1]]
        topics.append(", ".join(top_terms))
    return topics


# ------------------ Summarization Functions ------------------
abstractive_summarizer = get_abstractive_summarizer()


def abstractive_summarization(text):
    summary = abstractive_summarizer(text, max_length=50, min_length=25, do_sample=False)
    return summary[0]['summary_text']


def extractive_summarization(text, sentence_count=3):
    parser = PlaintextParser.from_string(text, Tokenizer("english"))
    summarizer = LexRankSummarizer()
    summary = summarizer(parser.document, sentence_count)
    return " ".join([str(sentence) for sentence in summary])


def query_based_summarization(text, query, threshold=0.1, top_n=2):
    sentences = sent_tokenize(text)
    if not sentences:
        return "No relevant information found for the query."
    action_indicators = ["urge", "request", "increase", "arrange", "immediate", "control", "measure", "action",
                         "implement", "improve", "take"]
    is_action_query = any(
        word in query.lower() for word in ["action", "request", "urge", "increase", "immediate", "control", "measure"])
    if is_action_query:
        threshold = 0.05
    corpus = sentences + [query]
    vectorizer_q = TfidfVectorizer().fit(corpus)
    sentence_vectors = vectorizer_q.transform(sentences)
    query_vector = vectorizer_q.transform([query])
    scores = np.dot(sentence_vectors, query_vector.T).toarray().flatten()
    valid_indices = [i for i, score in enumerate(scores) if score >= threshold]
    if not valid_indices:
        return "No relevant information found for the query."
    if is_action_query:
        valid_indices = [i for i in valid_indices if any(kw in sentences[i].lower() for kw in action_indicators)]
        if not valid_indices:
            return "No relevant information found for the query."
    sorted_indices = sorted(valid_indices, key=lambda i: scores[i], reverse=True)[:top_n]
    selected_indices = sorted(sorted_indices)
    summary = " ".join(sentences[i] for i in selected_indices)
    return summary


# ------------------ Personalized Summary Wrapper ------------------
def personalize_summary(summary, summary_type="general"):
    return f" {summary}"


# ------------------ Sentiment Analysis ------------------
sentiment_pipeline = get_sentiment_pipeline()


def sentiment_analysis(text):
    # Transformer Analysis
    transformer_result = sentiment_pipeline(text)[0]

    # Get highest confidence result
    sentiment_label = max(transformer_result, key=lambda x: x['score'])['label'] # "NEGATIVE" or "POSITIVE"
    confidence_score = max(transformer_result, key=lambda x: x['score'])['score']

    # Debug info is printed to console; remove or comment out if not needed.
    print("Transformer label:", sentiment_label)
    print("Transformer score:", confidence_score)

    # Generate explanation string
    explanation = f"Classified as {sentiment_label} with {confidence_score:.0%} confidence"

    return {
        "sentiment_label": sentiment_label,
        "confidence": confidence_score,
        "explanation": explanation
    }


# ------------------ Aggregated Analysis Plot Functions ------------------
def plot_classification_distribution(class_counts):
    fig = go.Figure([go.Bar(x=class_counts.index, y=class_counts.values)])
    fig.update_layout(title="Classification Distribution", xaxis_title="Category", yaxis_title="Count")
    return fig


def plot_sentiment_distribution(avg_sentiment):
    fig = go.Figure([go.Bar(x=["Average Sentiment Polarity"], y=[avg_sentiment])])
    fig.update_layout(title="Average Sentiment Polarity", xaxis_title="Metric", yaxis_title="Polarity")
    return fig


def plot_sentiment_gauge(confidence):
    fig = go.Figure(go.Indicator(
        mode="gauge+number",
        value=confidence,
        domain={'x': [0, 1], 'y': [0, 1]},
        title={"text": "Sentiment Confidence"},
        gauge={
            "axis": {"range": [0, 1]},
            "steps": [
                {"range": [0, 0.3], "color": "red"},
                {"range": [0.3, 0.7], "color": "yellow"},
                {"range": [0.7, 1], "color": "green"}
            ],
            "threshold": {
                "line": {"color": "black", "width": 4},
                "thickness": 0.75,
                "value": confidence
            }
        }
    ))
    fig.update_layout(
        annotations=[
            dict(x=0.15, y=0.1, text="<b>Low</b>", showarrow=False,
                 font=dict(color="red", size=12)),
            dict(x=0.50, y=0.1, text="<b>Medium</b>", showarrow=False,
                 font=dict(color="yellow", size=12)),
            dict(x=0.85, y=0.1, text="<b>High</b>", showarrow=False,
                 font=dict(color="green", size=12))
        ]
    )
    return fig



# ------------------ Report Generation Functions ------------------
from fpdf import FPDF
from docx import Document


def generate_pdf_report(original_text, abstractive_summary, extractive_summary, query_summary, sentiment_results):
    original_text = sanitize_text(original_text)
    abstractive_summary = sanitize_text(abstractive_summary)
    extractive_summary = sanitize_text(extractive_summary)
    query_summary = sanitize_text(query_summary)
    sentiment_results = sanitize_text(str(sentiment_results))
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(0, 10, txt="Analysis Report", ln=True, align='C')
    pdf.ln(5)
    pdf.multi_cell(0, 10, txt=f"Original Text:\n{original_text}\n")
    pdf.ln(3)
    pdf.multi_cell(0, 10, txt=f"Abstractive Summary:\n{abstractive_summary}\n")
    pdf.ln(3)
    pdf.multi_cell(0, 10, txt=f"Extractive Summary:\n{extractive_summary}\n")
    pdf.ln(3)
    pdf.multi_cell(0, 10, txt=f"Query-based Summary:\n{query_summary}\n")
    pdf.ln(3)
    pdf.multi_cell(0, 10, txt=f"Sentiment Analysis:\n{sentiment_results}\n")
    return pdf.output(dest='S').encode('latin1', errors='replace')


def generate_docx_report(original_text, abstractive_summary, extractive_summary, query_summary, sentiment_results):
    doc = Document()
    doc.add_heading("Analysis Report", level=1)
    doc.add_heading("Original Text", level=2)
    doc.add_paragraph(original_text)
    doc.add_heading("Abstractive Summary", level=2)
    doc.add_paragraph(abstractive_summary)
    doc.add_heading("Extractive Summary", level=2)
    doc.add_paragraph(extractive_summary)
    doc.add_heading("Query-based Summary", level=2)
    doc.add_paragraph(query_summary)
    doc.add_heading("Sentiment Analysis", level=2)
    doc.add_paragraph(str(sentiment_results))
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ------------------ Data Loading ------------------
def load_data(file_path):
    try:
        df = pd.read_csv(file_path, on_bad_lines='skip')
        st.write("Data loaded successfully.")
        return df
    except Exception as e:
        st.write("Error loading data:", e)
        return None


# ------------------ Main UI Pipeline Using Pages ------------------

def get_file_icon_path(num_files, file_extensions):
    """
    Return a path to the correct SVG icon
    given the number of files and the extension set.
    """
    if num_files == 0:
        return "src/img/Empty.svg"

    if num_files == 1:
        # Check if user pasted
        if "paste" in file_extensions:
            return "src/img/Single_default.svg"  # <--- For a single pasted letter
        elif "pdf" in file_extensions and len(file_extensions) == 1:
            return "src/img/Single_Pdf.svg"
        elif "doc" in file_extensions and len(file_extensions) == 1:
            return "src/img/Single_Doc.svg"
        else:
            return "src/img/Single_default.svg"

    # If multiple files
    # If all PDFs
    if file_extensions == {"pdf"}:
        return "src/img/Multiple_Pdf.svg"
    # If all docs
    elif file_extensions == {"doc"}:
        return "src/img/Multiple_Doc.svg"
    # If mixture
    elif "pdf" in file_extensions or "doc" in file_extensions:
        return "src/img/Multiple_Both.svg"
    else:
        return "src/img/Multiple_default.svg"

def pick_sidebar_icon(num_files, file_types):
    if num_files == 0:
        return "src/img/Multiple_Default.svg"
    if num_files == 1:
        ft = next(iter(file_types))
        if ft == "pdf":
            return "src/img/Single_Pdf.svg"
        elif ft == "doc":
            return "src/img/Single_Doc.svg"
        else:
            return "src/img/Single_Default.svg"
    # For multiple files
    if file_types == {"pdf"}:
        return "src/img/Multiple_Pdf.svg"
    elif file_types == {"doc"}:
        return "src/img/Multiple_Doc.svg"
    elif len(file_types) > 1:  # Mixed types
        return "src/img/Multiple_Both.svg"
    else:
        return "src/img/Multiple_Default.svg"

def data_entry_page():
    st.title("Letter Submission (Data Entry)")
    data_mode = st.radio("Choose Input Mode", ["Paste Text", "Upload File"])

    input_text = ""
    uploaded_files = []

    if data_mode == "Paste Text":
        input_text = st.text_area("Paste your letter text here", height=200)
    else:
        uploaded_files = st.file_uploader(
            "Upload files (txt, pdf, doc, docx)",
            type=["txt", "pdf", "doc", "docx"],
            accept_multiple_files=True
        )

    if st.button("Submit"):
        with st.spinner("Processing..."):
            if data_mode == "Paste Text":
                if not input_text.strip():
                    st.warning("Please paste some text before submitting.")
                    return

                st.session_state.input_text = input_text
                st.session_state.data_submitted = True
                st.session_state.data_mode = data_mode
                st.session_state.uploaded_file_info = {
                    "num_files": 1,
                    "file_extensions": {"paste"}
                }
                st.session_state.page = "results"
                st.rerun()

            elif data_mode == "Upload File":
                if not uploaded_files:
                    st.warning("Please upload at least one file.")
                    return

                file_types = []
                extracted_texts = []
                combined_text = ""

                for file in uploaded_files:
                    file_type = file.type
                    file_types.append(file_type)
                    text = ""

                    try:
                        if file_type == "application/pdf":
                            pdf_reader = PyPDF2.PdfReader(file)
                            text = "\n".join([
                                page.extract_text()
                                for page in pdf_reader.pages
                                if page.extract_text()
                            ])

                        elif file_type in ["application/msword",
                                         "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
                            doc = Document(file)
                            text = "\n".join([
                                para.text
                                for para in doc.paragraphs
                                if para.text.strip()
                            ])

                        elif file_type == "text/plain":
                            text = file.getvalue().decode("utf-8")

                    except Exception as e:
                        st.error(f"Error reading {file.name}: {str(e)}")
                        continue

                    if text.strip():
                        extracted_texts.append(text.strip())
                        combined_text += f"\n\n{text.strip()}"

                if not extracted_texts:
                    st.error("Could not extract any text from uploaded files")
                    return

                # Update session state
                st.session_state.uploaded_files_texts = extracted_texts
                st.session_state.input_text = combined_text.strip()
                st.session_state.data_submitted = True
                st.session_state.data_mode = data_mode

                # Handle file type icons (simplified to pdf/doc/other)
                ext_set = set()
                for ft in file_types:
                    ft_lower = ft.lower()
                    if "pdf" in ft_lower:
                        ext_set.add("pdf")
                    elif "msword" in ft_lower or "wordprocessingml.document" in ft_lower:
                        ext_set.add("doc")
                    else:
                        ext_set.add("other")

                st.session_state.uploaded_file_info = {
                    "num_files": len(uploaded_files),
                    "file_extensions": ext_set
                }

                # Route to correct page
                if len(uploaded_files) > 1:
                    st.session_state.page = "aggregated_analysis"
                else:
                    st.session_state.page = "results"

                st.rerun()


def results_page():
    st.title("Individual Letter Analysis")
    if "data_submitted" not in st.session_state or not st.session_state.data_submitted:
        st.warning("No data submitted yet. Please go to the 'Data Entry' page, provide a letter, and click Submit.")
        return

    # Get text based on input method
    if st.session_state.data_mode == "Upload File":
        if "uploaded_files_texts" in st.session_state and len(st.session_state.uploaded_files_texts) >= 1:
            letter_text = st.session_state.uploaded_files_texts[0]
        else:
            st.error("No text found in uploaded file")
            return
    else:
        letter_text = st.session_state.get("input_text", "")

    # Sidebar: Show a dynamic icon based on the upload type.
    # Sidebar: Unified icon handling
    with st.sidebar:
        file_info = st.session_state.get("uploaded_file_info", {})
        num_files = file_info.get("num_files", 0)
        ext_set = file_info.get("file_extensions", set())

        icon_path = pick_sidebar_icon(num_files, ext_set)
        st.image(icon_path, width=150)

        with st.expander("Original Letter", expanded=False):
            st.write(letter_text if letter_text.strip() else "No text available")

    # Continue processing for individual analysis
    letter_clean = comprehensive_text_preprocessing(letter_text)
    classifier = get_zero_shot_classifier()
    if letter_clean.strip():
        classification_result = classifier(letter_clean, candidate_labels)
        letter_class = classification_result["labels"][0]
    else:
        letter_class = "Unclassified (insufficient text)"
        st.warning("Could not classify - extracted text appears empty")
    st.subheader("Classification")
    st.write(f"This letter is classified as: **{letter_class}**")

    topic_label, top_keywords = compute_topic(letter_clean)
    st.subheader("Topic")
    st.write(f"Topic: **{topic_label}**")

    abstractive_res = abstractive_summarization(letter_text)
    extractive_res = extractive_summarization(letter_text)

    st.markdown(
        """
        <style>
        .card {
            box-shadow: 0 4px 8px 0 rgba(0,0,0,0.2);
            transition: 0.3s;
            border-radius: 5px;
            padding: 16px;
            margin: 10px 0;
            background-color: var(--background-color);
            color: var(--text-color);
        }
        .card:hover {
            box-shadow: 0 8px 16px 0 rgba(0,0,0,0.2);
        }
        [data-theme="light"] .card {
            --background-color: #F0F2F6;
            --text-color: #000000;
        }
        [data-theme="dark"] .card {
            --background-color: #262730;
            --text-color: #FFFFFF;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    col1, col2 = st.columns(2)
    with col1:
        st.subheader("💡 Key Takeaways")
        st.markdown(f"""
            <div class="card">
                <p>{abstractive_res}</p>
            </div>
        """, unsafe_allow_html=True)
    with col2:
        st.subheader("🔍 Highlighted Sentences")
        st.markdown(f"""
            <div class="card">
                <p>{extractive_res}</p>
            </div>
        """, unsafe_allow_html=True)

    st.subheader("❓ Inquisitive Summary")
    user_query = st.text_input("Ask anything about the letters:", "What actions are being urged in the letter?")
    query_res = query_based_summarization(letter_text, query=user_query)
    refined_query_res = paraphrase_text(query_res)
    st.write(personalize_summary(refined_query_res, "query"))
    # sentiment_label confidence
    st.subheader("🗣️ Resident Mood Overview")
    sentiment_results = sentiment_analysis(letter_text)
    sentiment_label = sentiment_results["sentiment_label"]
    confidence_score = sentiment_results["confidence"]
    explanation =  sentiment_results["explanation"]

    col_mood, col_gauge = st.columns(2)
    with col_mood:
        st.write(f"**Mood:** {sentiment_label}")
        st.write(explanation)
    with col_gauge:
        gauge_fig = plot_sentiment_gauge(confidence_score)
        st.plotly_chart(gauge_fig)

    export_format = st.selectbox("Select Export Format", ["PDF", "DOCX", "TXT", "CSV"])
    if export_format == "PDF":
        file_bytes = generate_pdf_report(letter_text, abstractive_res, extractive_res, query_res, sentiment_results)
        st.download_button("Download Report", file_bytes, file_name="analysis_report.pdf", mime="application/pdf")
    elif export_format == "DOCX":
        file_bytes = generate_docx_report(letter_text, abstractive_res, extractive_res, query_res, sentiment_results)
        st.download_button("Download Report", file_bytes, file_name="analysis_report.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    elif export_format == "TXT":
        txt_report = (
            f"Analysis Report\n\nOriginal Text:\n{letter_text}\n\nAbstractive Summary:\n{abstractive_res}\n\n"
            f"Extractive Summary:\n{extractive_res}\n\nQuery-based Summary:\n{query_res}\n\n"
            f"Sentiment Analysis:\n{sentiment_results}"
        )
        st.download_button("Download Report", txt_report, file_name="analysis_report.txt", mime="text/plain")
    elif export_format == "CSV":
        df_report = pd.DataFrame({
            "Original Text": [letter_text],
            "Abstractive Summary": [abstractive_res],
            "Extractive Summary": [extractive_res],
            "Query-based Summary": [query_res],
            "Sentiment Analysis": [str(sentiment_results)]
        })
        csv_report = df_report.to_csv(index=False)
        st.download_button("Download Report", csv_report, file_name="analysis_report.csv", mime="text/csv")

    if st.button("Back to Data Entry"):
        st.session_state.input_text = ""
        st.session_state.data_submitted = False
        st.session_state.page = "data_entry"
        st.rerun()







def aggregated_analysis_page():
    st.title("Comprehensive Letters Analysis")

    # 1) Check data availability
    if not st.session_state.get("data_submitted", False):
        st.warning("No data submitted yet. Please go to the 'Data Entry' page and upload multiple files.")
        return

    if "uploaded_files_texts" not in st.session_state or len(st.session_state.uploaded_files_texts) < 2:
        st.warning("No multiple-file data found. Please go to the 'Data Entry' page and upload multiple files.")
        return

    # 2) Sidebar content
    with st.sidebar:
        file_info = st.session_state.get("uploaded_file_info", {})
        num_files = file_info.get("num_files", 0)
        ext_set = file_info.get("file_extensions", set())

        icon_path = pick_sidebar_icon(num_files, ext_set)
        st.image(icon_path, width=150)

        with st.expander("Original Letter", expanded=False):
            st.write(st.session_state.get("input_text", "No text available."))

    # 3) Prepare data
    uploaded_texts = st.session_state.uploaded_files_texts
    df_agg = pd.DataFrame({"text": uploaded_texts})
    df_agg["clean_text"] = df_agg["text"].apply(comprehensive_text_preprocessing)
    texts_clean = df_agg["clean_text"].tolist()

    # 4) Extract addresses, issues, and topics from the text
    df_agg["Address"] = df_agg["text"].apply(lambda x: extract_locations(x))
    df_agg["Issue"], df_agg["Topic"] = zip(*df_agg["text"].apply(lambda x: compute_topic(x)))
    df_agg["sentiment"] = df_agg["text"].apply(lambda x: sentiment_analysis(x)["sentiment_label"])
    df_agg["classification"] = df_agg["clean_text"].apply(classify_document)

    # 5) Print the extracted data to the console
    print("\nExtracted Data for Map Visualization:")
    for index, row in df_agg.iterrows():
        address = row["Address"] if row["Address"] else "No Address Extracted"
        topic = row["Topic"] if row["Topic"] else "No Topic Extracted"
        sentiment = row["sentiment"] if row["sentiment"] else "No Sentiment Extracted"
        issue = row["Issue"] if row["Issue"] else "No Issue Extracted"
        category = row["classification"] if row["classification"] else "No Category Extracted"

        print(f"\nLetter {index + 1}:")
        print(f"  - Address: {address}")
        print(f"  - Topic: {topic}")
        print(f"  - Sentiment: {sentiment}")
        print(f"  - Issue: {issue}")
        print(f"  - Category: {category}")

    # 6) Geocode addresses
    df_agg["Address"] = df_agg["text"].apply(lambda x: extract_locations(x))
    lat_lon_list = geocode_addresses(df_agg["Address"].tolist())
    df_agg[['lat', 'lon']] = pd.DataFrame(lat_lon_list, columns=['lat', 'lon'])

    # 7) Calculate Key Metrics
    total_letters = len(df_agg)
    class_counts = df_agg["classification"].value_counts(normalize=True) * 100
    local_problems_pct = class_counts.get("Local Problem", 0)
    new_initiatives_pct = class_counts.get("New Initiatives", 0)

    # Key Metrics Cards
    st.markdown("### Key Metrics")
    kpi_col1, kpi_col2, kpi_col3 = st.columns(3)

    theme_base = st.get_option("theme.base")
    bg_color = "#FFFFFF" if theme_base == "light" else "#222"
    text_color = "#000000" if theme_base == "light" else "#FFFFFF"

    with kpi_col1:
        st.markdown(f"""
        <div style='background-color:{bg_color}; padding:15px; border-radius:10px; text-align:center; box-shadow: 0px 4px 6px rgba(0,0,0,0.1);'>
            <h3 style='color:{text_color};'>📩 Total Letters</h3>
            <h2 style='color:{text_color};'>{total_letters}</h2>
        </div>
        """, unsafe_allow_html=True)

    with kpi_col2:
        st.markdown(f"""
        <div style='background-color:{bg_color}; padding:15px; border-radius:10px; text-align:center; box-shadow: 0px 4px 6px rgba(0,0,0,0.1);'>
            <h3 style='color:{text_color};'>📍 Local Problems</h3>
            <h2 style='color:{text_color};'>{local_problems_pct:.1f}%</h2>
        </div>
        """, unsafe_allow_html=True)

    with kpi_col3:
        st.markdown(f"""
        <div style='background-color:{bg_color}; padding:15px; border-radius:10px; text-align:center; box-shadow: 0px 4px 6px rgba(0,0,0,0.1);'>
            <h3 style='color:{text_color};'>✨ New Initiatives</h3>
            <h2 style='color:{text_color};'>{new_initiatives_pct:.1f}%</h2>
        </div>
        """, unsafe_allow_html=True)

    # 8) Most Common Issues
    st.subheader("Most Common Issues")

    # Get topics and their keywords
    topics = topic_modeling(texts_clean, num_topics=3)

    # Create regex patterns for exact word matching
    issues_data = []
    for topic in topics:
        keywords = [re.escape(kw.strip()) for kw in topic.split(',')]
        pattern = r'\b(' + '|'.join(keywords) + r')\b'

        # Count matches with case insensitivity
        count = df_agg['clean_text'].str.contains(
            pattern,
            regex=True,
            case=False,
            na=False
        ).sum()

        issues_data.append({
            "Issue": dynamic_topic_label(topic),
            "Count": count,
            "Percentage": (count / len(df_agg) * 100)
        })

    # Create and sort DataFrame
    issues_df = pd.DataFrame(issues_data).sort_values('Count', ascending=False)

    # Format percentages
    issues_df['Percentage'] = issues_df['Percentage'].round(1)

    # Create visualization with combined labels
    fig = px.bar(
        issues_df,
        x="Issue",
        y="Count",
        text="Percentage",
        labels={'Count': 'Number of Complaints', 'Percentage': 'Percentage'},
        color="Issue"
    )

    # Improve label formatting
    fig.update_traces(
        texttemplate='%{text}%',
        textposition='outside',
        hovertemplate="<b>%{x}</b><br>Count: %{y}<br>Percentage: %{text}%"
    )

    # Set axis limits based on data
    max_count = issues_df['Count'].max()
    fig.update_layout(
        yaxis_range=[0, max_count * 1.2],
        showlegend=False
    )

    st.plotly_chart(fig, use_container_width=True)

    # 9) Classification & Sentiment Analysis
    st.subheader("📊 Classification Distribution & 😊 Sentiment Analysis")
    col4, col5 = st.columns(2)

    with col4:
        # Classification Distribution
        class_counts = df_agg["classification"].value_counts()
        fig_classification = px.pie(
            class_counts,
            values=class_counts.values,
            names=class_counts.index,
            title="Classification Distribution"
        )
        st.plotly_chart(fig_classification, use_container_width=True)

    with col5:
        # Sentiment Analysis
        sentiment_counts = df_agg["sentiment"].value_counts()
        fig_sentiment = px.bar(
            sentiment_counts,
            x=sentiment_counts.index,
            y=sentiment_counts.values,
            title="Sentiment Analysis",
            color=sentiment_counts.index
        )
        st.plotly_chart(fig_sentiment, use_container_width=True)

    # 10) Key Takeaways & Highlighted Sentences
    col6, col7 = st.columns(2)

    with col6:
        st.subheader("💡 Key Takeaways")
        try:
            key_takeaways = " ".join([abstractive_summarization(text)
                                      for text in st.session_state.uploaded_files_texts[:3]])
            st.markdown(f"""
                <div class="card">
                    <p>{key_takeaways[:500]}</p>  
                </div>
            """, unsafe_allow_html=True)
        except Exception as e:
            st.error(f"Key Takeaways Error: {str(e)}")
            st.session_state.page = "data_entry"
            st.rerun()

    with col7:
        st.subheader("🔍 Highlighted Sentences")
        try:
            highlighted = " ".join([extractive_summarization(text, 1)
                                    for text in st.session_state.uploaded_files_texts[:3]])
            st.markdown(f"""
                <div class="card">
                    <p>{highlighted[:500]}</p>  
                </div>
            """, unsafe_allow_html=True)
        except Exception as e:
            st.error(f"Highlight Error: {str(e)}")
            st.session_state.page = "data_entry"
            st.rerun()

    # AI Search Section
    st.subheader("🔍 AI Document Analyst")
    user_question = st.text_input("Ask anything about the letters:",
                                  placeholder="e.g. What are the main complaints about waste management?")

    if user_question:
        with st.spinner("Analyzing documents..."):
            try:
                # Get all uploaded texts
                documents = st.session_state.uploaded_files_texts

                # Get AI response
                response = ai_question_answer(user_question, documents)

                # Display response
                st.markdown(f"""
                <div style='
                    padding: 15px;
                    border-radius: 10px;
                    background-color: #f0f2f6;
                    margin: 10px 0;
                    box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                '>
                    <p style='font-size: 16px; color: #333;'>{response}</p>
                </div>
                """, unsafe_allow_html=True)

            except Exception as e:
                st.error(f"Error processing question: {str(e)}")

    # 11) Create the map with tabs
    st.subheader("📍 Geographic Issue Distribution")



    # Create tabs
    tab1, tab2, tab3, tab4 = st.tabs(["Show Map by Sentiment", "Show Map by Categories", "Show Map by Topics", "Show Map by Issues"])

    with tab1:
        st.write("### Map Filtered by Sentiment")
        if df_agg.empty or df_agg["lat"].isnull().all() or df_agg["lon"].isnull().all():
            st.warning("No geographic data available for mapping.")
        else:
            # Map for Sentiment
            df_sentiment = df_agg.dropna(subset=["lat", "lon", "sentiment"])
            if df_sentiment.empty:
                st.warning("No data found for sentiment mapping.")
            else:
                # Color points based on sentiment
                df_sentiment["color"] = df_sentiment["sentiment"].apply(
                    lambda x: [0, 255, 0, 160] if x == "POSITIVE" else [255, 0, 0, 160])

                # Calculate the bounding box of the data points
                min_lat, max_lat = df_sentiment["lat"].min(), df_sentiment["lat"].max()
                min_lon, max_lon = df_sentiment["lon"].min(), df_sentiment["lon"].max()

                # Create PyDeck map
                layer = pdk.Layer(
                    "ScatterplotLayer",
                    data=df_sentiment,
                    get_position=["lon", "lat"],
                    get_color="color",
                    get_radius=200,
                    pickable=True,
                )
                view_state = pdk.ViewState(
                    latitude=(min_lat + max_lat) / 2,
                    longitude=(min_lon + max_lon) / 2,
                    zoom=10,
                    min_zoom=5,
                    max_zoom=15,
                    pitch=0,
                    bearing=0
                )
                deck = pdk.Deck(
                    layers=[layer],
                    initial_view_state=view_state,
                    map_style="mapbox://styles/mapbox/light-v10" if st.get_option(
                        "theme.base") == "light" else "mapbox://styles/mapbox/dark-v10",
                    tooltip={"text": "Sentiment: {sentiment}\nAddress: {Address}"}
                )
                st.pydeck_chart(deck)

    with tab2:
        st.write("### Map Filtered by Categories")
        if df_agg.empty or df_agg["lat"].isnull().all() or df_agg["lon"].isnull().all():
            st.warning("No geographic data available for mapping.")
        else:
            # Map for Categories
            df_category = df_agg.dropna(subset=["lat", "lon", "classification"])
            if df_category.empty:
                st.warning("No data found for category mapping.")
            else:
                # Assign unique colors to each category
                unique_categories = df_category["classification"].unique()
                base_colors = px.colors.qualitative.Plotly
                if len(unique_categories) > len(base_colors):
                    # If more categories than colors, cycle through the base colors
                    color_map = {category: base_colors[i % len(base_colors)] for i, category in
                                 enumerate(unique_categories)}
                else:
                    # Otherwise, use the base colors
                    color_map = {category: base_colors[i] for i, category in enumerate(unique_categories)}
                df_category["color"] = df_category["classification"].map(color_map)

                # Calculate the bounding box of the data points
                min_lat, max_lat = df_category["lat"].min(), df_category["lat"].max()
                min_lon, max_lon = df_category["lon"].min(), df_category["lon"].max()

                # Create PyDeck map
                layer = pdk.Layer(
                    "ScatterplotLayer",
                    data=df_category,
                    get_position=["lon", "lat"],
                    get_color="color",
                    get_radius=200,
                    pickable=True,
                )
                view_state = pdk.ViewState(
                    latitude=(min_lat + max_lat) / 2,
                    longitude=(min_lon + max_lon) / 2,
                    zoom=10,
                    min_zoom=5,
                    max_zoom=15,
                    pitch=0,
                    bearing=0
                )
                deck = pdk.Deck(
                    layers=[layer],
                    initial_view_state=view_state,
                    map_style="mapbox://styles/mapbox/light-v10" if st.get_option(
                        "theme.base") == "light" else "mapbox://styles/mapbox/dark-v10",
                    tooltip={"text": "Category: {classification}\nAddress: {Address}"}
                )
                st.pydeck_chart(deck)

    with tab3:
        st.write("### Map Filtered by Topics")
        if df_agg.empty or df_agg["lat"].isnull().all() or df_agg["lon"].isnull().all():
            st.warning("No geographic data available for mapping.")
        else:
            # Map for Topics
            df_topic = df_agg.dropna(subset=["lat", "lon", "Topic"])
            if df_topic.empty:
                st.warning("No data found for topic mapping.")
            else:
                # Assign unique colors to each topic
                unique_topics = df_topic["Topic"].unique()
                base_colors = px.colors.qualitative.Plotly
                if len(unique_topics) > len(base_colors):
                    # If more topics than colors, cycle through the base colors
                    color_map = {topic: base_colors[i % len(base_colors)] for i, topic in enumerate(unique_topics)}
                else:
                    # Otherwise, use the base colors
                    color_map = {topic: base_colors[i] for i, topic in enumerate(unique_topics)}
                df_topic["color"] = df_topic["Topic"].map(color_map)

                # Calculate the bounding box of the data points
                min_lat, max_lat = df_topic["lat"].min(), df_topic["lat"].max()
                min_lon, max_lon = df_topic["lon"].min(), df_topic["lon"].max()

                # Create PyDeck map
                layer = pdk.Layer(
                    "ScatterplotLayer",
                    data=df_topic,
                    get_position=["lon", "lat"],
                    get_color="color",
                    get_radius=200,
                    pickable=True,
                )
                view_state = pdk.ViewState(
                    latitude=(min_lat + max_lat) / 2,
                    longitude=(min_lon + max_lon) / 2,
                    zoom=10,
                    min_zoom=5,
                    max_zoom=15,
                    pitch=0,
                    bearing=0
                )
                deck = pdk.Deck(
                    layers=[layer],
                    initial_view_state=view_state,
                    map_style="mapbox://styles/mapbox/light-v10" if st.get_option(
                        "theme.base") == "light" else "mapbox://styles/mapbox/dark-v10",
                    tooltip={"text": "Topic: {Topic}\nAddress: {Address}"}
                )
                st.pydeck_chart(deck)

    with tab4:
        st.write("### Map Filtered by Issues")
        if df_agg.empty or df_agg["lat"].isnull().all() or df_agg["lon"].isnull().all():
            st.warning("No geographic data available for mapping.")
        else:
            # Map for Issues
            df_issue = df_agg.dropna(subset=["lat", "lon", "Issue"])
            if df_issue.empty:
                st.warning("No data found for issue mapping.")
            else:
                # Assign unique colors to each issue
                unique_issues = df_issue["Issue"].unique()
                base_colors = px.colors.qualitative.Plotly
                if len(unique_issues) > len(base_colors):
                    # If more issues than colors, cycle through the base colors
                    color_map = {issue: base_colors[i % len(base_colors)] for i, issue in enumerate(unique_issues)}
                else:
                    # Otherwise, use the base colors
                    color_map = {issue: base_colors[i] for i, issue in enumerate(unique_issues)}
                df_issue["color"] = df_issue["Issue"].map(color_map)

                # Calculate the bounding box of the data points
                min_lat, max_lat = df_issue["lat"].min(), df_issue["lat"].max()
                min_lon, max_lon = df_issue["lon"].min(), df_issue["lon"].max()

                # Create PyDeck map
                layer = pdk.Layer(
                    "ScatterplotLayer",
                    data=df_issue,
                    get_position=["lon", "lat"],
                    get_color="color",
                    get_radius=200,
                    pickable=True,
                )
                view_state = pdk.ViewState(
                    latitude=(min_lat + max_lat) / 2,
                    longitude=(min_lon + max_lon) / 2,
                    zoom=10,
                    min_zoom=5,
                    max_zoom=15,
                    pitch=0,
                    bearing=0
                )
                deck = pdk.Deck(
                    layers=[layer],
                    initial_view_state=view_state,
                    map_style="mapbox://styles/mapbox/light-v10" if st.get_option(
                        "theme.base") == "light" else "mapbox://styles/mapbox/dark-v10",
                    tooltip={"text": "Issue: {Issue}\nAddress: {Address}"}
                )
                st.pydeck_chart(deck)

    # 12) Export options
    st.subheader("Export Options")
    report_csv = df_agg.to_csv(index=False)
    st.download_button("Download Report (CSV)", report_csv, file_name="aggregated_report.csv", mime="text/csv")

    # 13) Navigation
    if st.button("Back to Data Entry"):
        st.session_state.input_text = ""
        st.session_state.data_submitted = False
        st.session_state.page = "data_entry"
        st.rerun()



def main():
    st.set_page_config(layout="wide", page_title="Bolsover District Council - Analysis")
    st.sidebar.image("src/img/Bolsover_District_Council_logo.svg", width=150)

    # Only if data has been submitted in Upload File mode, show dynamic icon and original letter expander

    # Page selection logic
    if "page" not in st.session_state:
        st.session_state.page = "data_entry"

    if st.session_state.page == "data_entry":
        data_entry_page()
    elif st.session_state.page == "results":
        results_page()
    elif st.session_state.page == "aggregated_analysis":
        aggregated_analysis_page()


if __name__ == '__main__':
    main()
