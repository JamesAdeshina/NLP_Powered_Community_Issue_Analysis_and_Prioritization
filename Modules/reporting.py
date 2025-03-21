import io
from fpdf import FPDF
from docx import Document
from preprocessing import sanitize_text

def generate_pdf_report(original_text: str, abstractive_summary: str, extractive_summary: str, query_summary: str, sentiment_results) -> bytes:
    original_text = sanitize_text(original_text)
    abstractive_summary = sanitize_text(abstractive_summary)
    extractive_summary = sanitize_text(extractive_summary)
    query_summary = sanitize_text(query_summary)
    sentiment_results = sanitize_text(str(sentiment_results))
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(0, 10, txt="Analysis Report", ln=True, align='C')
    pdf.ln(5)
    pdf.multi_cell(0, 10, txt=f"Original Text:\n{original_text}\n")
    pdf.ln(3)
    pdf.multi_cell(0, 10, txt=f"Abstractive Summary:\n{abstractive_summary}\n")
    pdf.ln(3)
    pdf.multi_cell(0, 10, txt=f"Extractive Summary:\n{extractive_summary}\n")
    pdf.ln(3)
    pdf.multi_cell(0, 10, txt=f"Query-based Summary:\n{query_summary}\n")
    pdf.ln(3)
    pdf.multi_cell(0, 10, txt=f"Sentiment Analysis:\n{sentiment_results}\n")
    return pdf.output(dest='S').encode('latin1', errors='replace')

def generate_docx_report(original_text: str, abstractive_summary: str, extractive_summary: str, query_summary: str, sentiment_results) -> bytes:
    doc = Document()
    doc.add_heading("Analysis Report", level=1)
    doc.add_heading("Original Text", level=2)
    doc.add_paragraph(original_text)
    doc.add_heading("Abstractive Summary", level=2)
    doc.add_paragraph(abstractive_summary)
    doc.add_heading("Extractive Summary", level=2)
    doc.add_paragraph(extractive_summary)
    doc.add_heading("Query-based Summary", level=2)
    doc.add_paragraph(query_summary)
    doc.add_heading("Sentiment Analysis", level=2)
    doc.add_paragraph(str(sentiment_results))
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
